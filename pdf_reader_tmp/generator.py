import os
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def set_cell_margins(cell, **kwargs):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m in [
        "top",
        "start",
        "bottom",
        "end",
    ]:
        if m in kwargs:
            node = OxmlElement("w:{}".format(m))
            node.set(qn('w:w'), str(kwargs.get(m)))
            node.set(qn('w:type'), 'dxa')
            tcMar.append(node)
    tcPr.append(tcMar)

doc = Document()

# Set global default font to Times New Roman, Size 12
style = doc.styles['Normal']
font = style.font
font.name = 'Times New Roman'
font.size = Pt(12)
font.color.rgb = RGBColor(0, 0, 0)

# Set global paragraph spacing
style.paragraph_format.line_spacing = 1.5
style.paragraph_format.space_after = Pt(6)

def add_heading_styled(text, size=16, align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, underline=False):
    p = doc.add_paragraph()
    p.alignment = align
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(size)
    run.bold = bold
    run.underline = underline
    run.font.color.rgb = RGBColor(0, 0, 0)
    return p

def add_para(text, align=WD_ALIGN_PARAGRAPH.JUSTIFY, bold=False, size=12):
    p = doc.add_paragraph()
    p.alignment = align
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(size)
    run.bold = bold
    run.font.color.rgb = RGBColor(0, 0, 0)
    return p

def add_space(lines=1):
    for _ in range(lines):
        doc.add_paragraph()

# ----------------- COVER PAGE -----------------
add_heading_styled("SUSHGANGA POLYTECHNIC, WANI", 20, WD_ALIGN_PARAGRAPH.CENTER, True)
add_space(2)
add_heading_styled("PROJECT REPORT", 16, WD_ALIGN_PARAGRAPH.CENTER, True)
add_heading_styled("ON", 14, WD_ALIGN_PARAGRAPH.CENTER, True)
add_space(1)
add_heading_styled("CARGO", 26, WD_ALIGN_PARAGRAPH.CENTER, True)
add_space(2)

table = doc.add_table(rows=1, cols=2)
table.autofit = False
table.allow_autofit = False
cell_left = table.cell(0, 0)
cell_right = table.cell(0, 1)

p = cell_left.paragraphs[0]
p.alignment = WD_ALIGN_PARAGRAPH.LEFT
r = p.add_run("SUBMITTED BY:")
r.bold = True
r.font.size = Pt(14)
r.font.name = 'Times New Roman'
p.add_run("\nRohan Rathod\nAditya Khaire\nAyush Kambale").font.size = Pt(14)

p = cell_right.paragraphs[0]
p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
r = p.add_run("GUIDED BY:")
r.bold = True
r.font.size = Pt(14)
r.font.name = 'Times New Roman'
p.add_run("\nMiss Diksha Hiware\n").font.size = Pt(14)
r2 = p.add_run("\nH.O.D:")
r2.bold = True
r2.font.size = Pt(14)
p.add_run("\nMrs. Shamli Kadu").font.size = Pt(14)

add_space(4)
add_heading_styled("DEPARTMENT OF COMPUTER TECHNOLOGY", 16, WD_ALIGN_PARAGRAPH.CENTER, True)
add_heading_styled("DIPLOMA IN COMPUTER TECHNOLOGY", 14, WD_ALIGN_PARAGRAPH.CENTER, True)
add_heading_styled("SUSHGANGA POLYTECHNIC, WANI", 16, WD_ALIGN_PARAGRAPH.CENTER, True)
add_heading_styled("SESSION 2024-2025", 14, WD_ALIGN_PARAGRAPH.CENTER, True)
doc.add_page_break()

# ----------------- CERTIFICATE -----------------
add_heading_styled("Certificate", 20, WD_ALIGN_PARAGRAPH.CENTER, True, True)
add_space(1)
add_heading_styled("THIS IS TO CERTIFY THAT PROJECT ON:", 14, WD_ALIGN_PARAGRAPH.CENTER, False)
add_heading_styled("CARGO", 18, WD_ALIGN_PARAGRAPH.CENTER, True)
add_space(1)
add_heading_styled("SUBMITTED BY", 14, WD_ALIGN_PARAGRAPH.CENTER, True)
add_heading_styled("ROHAN RATHOD", 12, WD_ALIGN_PARAGRAPH.CENTER, True)
add_heading_styled("ADITYA KHAIRE", 12, WD_ALIGN_PARAGRAPH.CENTER, True)
add_heading_styled("AYUSH KAMBALE", 12, WD_ALIGN_PARAGRAPH.CENTER, True)
add_space(1)

cert_text = "The Students of final year DIPLOMA IN COMPUTER TECHNOLOGY of Institute SUSHGANGA POLYTECHNIC, has successfully completed the project work and have submitted satisfactory report as guided by Miss Diksha Hiware on topic as per the syllabus prescribed by MAHARASHTRA STATE BOARD OF TECHNICAL EDUCATION, MUMBAI during academic session 2024-25."
add_para(cert_text, WD_ALIGN_PARAGRAPH.JUSTIFY)
add_space(3)

sig_table = doc.add_table(rows=2, cols=4)
sig_table.autofit = True

cell1 = sig_table.cell(0, 0)
cell1.text = "____________________\nMiss Diksha Hiware\nProject Guide"

cell2 = sig_table.cell(0, 1)
cell2.text = "____________________\nMrs. Shamli Kadu\nH.O.D C.M. dept"

cell3 = sig_table.cell(0, 2)
cell3.text = "____________________\nPrincipal\nS.G.P. Wani"

cell4 = sig_table.cell(0, 3)
cell4.text = "____________________\nExternal\n"

for row in sig_table.rows:
    for cell in row.cells:
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.font.name = 'Times New Roman'
                run.font.size = Pt(12)
                run.font.bold = True

doc.add_page_break()

# ----------------- SUBMISSION -----------------
add_heading_styled("SUBMISSION", 18, WD_ALIGN_PARAGRAPH.CENTER, True, True)
add_space(1)
sub_p1 = "We, Students of final year of course computer technology humbly submit that we have completed from time to time project works or described in this report by our own skills and study between period 2024-25 as per instructions and guidance of our guide."
sub_p2 = "And that, following students were associated with us in this work however quantum of our contribution has been approved by our lecturer. We conclude that we have not copied the report or its any appreciable part from other literature in contravention of academic ethics."
add_para(sub_p1)
add_para(sub_p2)
add_space(2)

date_p = doc.add_paragraph()
r = date_p.add_run("Date: __/__/2025")
r.font.name = 'Times New Roman'
r.font.size = Pt(12)

add_space(1)
add_para("Signature of students with respective names:", WD_ALIGN_PARAGRAPH.LEFT, True)
add_para("Rohan Rathod       ______________", WD_ALIGN_PARAGRAPH.LEFT)
add_para("Aditya Khaire      ______________", WD_ALIGN_PARAGRAPH.LEFT)
add_para("Ayush Kambale      ______________", WD_ALIGN_PARAGRAPH.LEFT)
doc.add_page_break()

# ----------------- ACKNOWLEDGEMENT -----------------
add_heading_styled("Acknowledgement", 18, WD_ALIGN_PARAGRAPH.CENTER, True, True)
add_space(1)
add_para("We sincerely thank Miss Diksha Hiware for her guidance and support, which played a vital role in this project.")
add_para("A big thanks to my team members – Rohan, Aditya, and Ayush, for their support and contribution.")
doc.add_page_break()

# ----------------- CONTENTS -----------------
add_heading_styled("Contents", 18, WD_ALIGN_PARAGRAPH.CENTER, True, True)
add_space(1)

contents_table = doc.add_table(rows=1, cols=2)
contents_table.autofit = True
hdr_cells = contents_table.rows[0].cells
hdr_cells[0].text = 'Sr. No'
hdr_cells[1].text = 'Content'

for cell in hdr_cells:
    for p in cell.paragraphs:
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        for r in p.runs:
            r.font.bold = True
            r.font.size = Pt(14)
            r.font.name = 'Times New Roman'

chapters = [
    "Introduction", "Objective", "Logo", "Preferred IDE & Hardware Used",
    "Technology & Language Used", "Home Screen / Landing Page", "Login Screen",
    "Registration Screen", "Admin Dashboard", "Booking Screen",
    "Profile Screen", "Future Development", "Bibliography"
]

for i, chapter in enumerate(chapters, 1):
    row_cells = contents_table.add_row().cells
    row_cells[0].text = str(i)
    row_cells[1].text = chapter
    for cell in row_cells:
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            for r in p.runs:
                r.font.size = Pt(12)
                r.font.name = 'Times New Roman'

doc.add_page_break()

# ----------------- CHAPTERS -----------------
def add_chapter(num, title, text):
    add_heading_styled(f"{num}. {title}", 16, WD_ALIGN_PARAGRAPH.LEFT, True)
    add_space(1)
    if isinstance(text, list):
        for line in text:
            if line.startswith("CODE:"):
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                r = p.add_run(line.replace("CODE:", ""))
                r.font.name = 'Courier New'
                r.font.size = Pt(10)
            else:
                add_para(line)
    else:
        add_para(text)
    doc.add_page_break()

add_chapter(1, "Introduction", "Cargo is a smart transport aggregator platform. Today, platforms like Ola and Uber connect passengers with drivers. But they mainly work in big cities and use fixed pricing. Cargo is different because it connects passengers, drivers, and vehicle owners directly. Our platform allows drivers and owners to set their own travel prices which is very useful for rural and non-metro areas. It also has a unique vehicle rental system where people who do not own a vehicle can rent one and work.")

add_chapter(2, "Objective", ["• To create a simple and easy to use transport app for booking rides and sending goods.", "• To let drivers decide their own price per km instead of the company deciding it.", "• To help vehicle owners rent out their unused vehicles to earn money.", "• To provide a single platform connecting all transport services."])

add_chapter(3, "Logo", "[ Paste Cargo Logo Image Here ]")

add_chapter(4, "Preferred IDE & Hardware Used", ["IDE: Visual Studio Code (VS Code).", "It is a very fast and popular code editor. We used VS Code to write all our Next.js and Tailwind code. It has great extensions for auto-complete.", "", "Hardware Used:", "• Laptop: Windows 10/11 operating system.", "• Minimum 8GB RAM to run local servers smoothly."])

add_chapter(5, "Technology & Language Used", ["Frontend:", "• Next.js 15: This is a fast React framework we used to make the website quickly.", "• Tailwind CSS: We used this to style the pages beautifully.", "• TypeScript: We used this to reduce errors in code.", "", "Backend & Database:", "• Node.js & Next.js API Routes: For the server-side logic.", "• MongoDB: It is a NoSQL database which is easy to store user details and booking states. We used Mongoose to connect to it.", "• Cloudinary: For saving user profile images."])

add_chapter(6, "Home Screen / Landing Page", ["[ Paste Landing Page Screenshot Here ]", "", "Code Snippet (Navbar.tsx):", "CODE:export default function Navbar() {\n  return (\n    <nav className=\"flex justify-between p-4 bg-white\">\n      <div>Cargo Logo</div>\n      <div className=\"flex gap-4\">\n        <Link href=\"/login\">Login</Link>\n        <Link href=\"/signup\">Sign Up</Link>\n      </div>\n    </nav>\n  )\n}"])

add_chapter(7, "Login Screen", ["[ Paste Login Screenshot Here ]", "", "Code Snippet (login/page.tsx):", "CODE:export default function Login() {\n  return (\n    <div className=\"flex items-center justify-center min-h-screen\">\n      <form className=\"p-6 bg-white shadow-md\">\n         <h2>Login to Cargo</h2>\n         <input type=\"email\" placeholder=\"Email\" />\n         <input type=\"password\" placeholder=\"Password\" />\n         <button>Login</button>\n      </form>\n    </div>\n  )\n}"])

add_chapter(8, "Registration Screen", ["[ Paste Registration Screenshot Here ]", "", "Code Snippet (signup/page.tsx):", "CODE:export default function Register() {\n  return (\n    <form className=\"flex flex-col gap-4\">\n      <h2>Register Account</h2>\n      <select>\n         <option>Passenger</option>\n         <option>Driver</option>\n         <option>Owner</option>\n      </select>\n      <input type=\"text\" placeholder=\"Name\" />\n      <input type=\"email\" placeholder=\"Email\" />\n      <button>Sign Up</button>\n    </form>\n  )\n}"])

add_chapter(9, "Admin Dashboard", ["[ Paste Admin Dashboard Screenshot Here ]", "", "In this page, the admin can manage all users, bookings, and system settings."])

add_chapter(10, "Booking Screen", ["[ Paste Find Ride Area Screenshot Here ]", "", "This is where passengers can select their pickup and drop locations. They can directly see drivers and their custom rates."])

add_chapter(11, "Profile Screen", ["[ Paste User Profile Screenshot Here ]", "", "Users can upload their license, update their name and manage their cars from the profile."])

add_chapter(12, "Future Development", ["In the future, we plan to add:", "• Real-time live tracking of vehicles using GPS.", "• Online payment like Google Pay and PhonePe using Razorpay Gateway.", "• Direct messaging chat system inside the app to let driver and passenger talk.", "• A specific Mobile App (Android/iOS) using React Native."])

add_chapter(13, "Bibliography", ["1. Next.js Documentation (nextjs.org)", "2. MongoDB Manual (mongodb.com)", "3. Tailwind CSS Docs (tailwindcss.com)", "4. Vercel for platform hosting", "5. StackOverflow and YouTube tutorials for debugging."])

# Save
path = r"c:\Users\telan\OneDrive\Desktop\cargo\Cargo_Thesis_V2.docx"
doc.save(path)
print(f"Generated successfully: {path}")
